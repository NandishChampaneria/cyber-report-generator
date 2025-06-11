from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import os
import glob
from pathlib import Path

def get_organization_info(logo_folder_path):
    """Get organization name and logo path from logo folder or direct file path"""
    org_name = None
    org_logo_path = None
    
    if not logo_folder_path:
        print("📄 No logo path provided")
        return org_name, org_logo_path
    
    # Check if the path is a direct file path or a folder path
    if os.path.isfile(logo_folder_path):
        # Direct file path provided
        org_logo_path = logo_folder_path
        # Extract organization name from filename (remove extension and path)
        filename = Path(org_logo_path).stem
        # Convert filename to proper organization name (capitalize words, replace underscores/hyphens with spaces)
        org_name = filename.replace('_', ' ').replace('-', ' ').title()
        print(f"📄 Found organization logo: {org_logo_path}")
        print(f"📄 Organization name: {org_name}")
    elif os.path.isdir(logo_folder_path):
        # Folder path provided - look for organization logos (excluding sequretek.png)
        logo_files = []
        # Get all image files except sequretek.png
        for ext in ['*.png', '*.jpg', '*.jpeg']:
            files = glob.glob(os.path.join(logo_folder_path, ext))
            logo_files.extend([f for f in files if not f.endswith('sequretek.png')])
        
        if logo_files:
            # Use the first organization logo found
            org_logo_path = logo_files[0]
            # Extract organization name from filename (remove extension and path)
            filename = Path(org_logo_path).stem
            # Convert filename to proper organization name (capitalize words, replace underscores/hyphens with spaces)
            org_name = filename.replace('_', ' ').replace('-', ' ').title()
            print(f"📄 Found organization logo: {org_logo_path}")
            print(f"📄 Organization name: {org_name}")
        else:
            print("📄 No organization logo found in logo folder (other than sequretek.png)")
    else:
        print(f"📄 Logo path does not exist: {logo_folder_path}")
    
    return org_name, org_logo_path

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    
    # Add bookmark for TOC reference
    bookmark_name = text.replace(' ', '_').replace('&', 'and').replace(':', '').replace('/', '')
    
    # Generate a unique ID for the bookmark
    unique_id = str(abs(hash(bookmark_name)) % 999999999)
    
    # Add bookmark to the heading for page reference
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), unique_id)
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), unique_id)
    
    # Insert bookmark around the heading text
    heading_p = heading._element
    heading_p.insert(0, bookmark_start)
    heading_p.append(bookmark_end)
    
    return heading

def add_paragraph(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_image(doc, image_path, width_in_inches=6):
    doc.add_picture(image_path, width=Inches(width_in_inches))

def add_table(doc, data, column_names):
    table = doc.add_table(rows=1, cols=len(column_names))
    table.style = 'Table Grid'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(column_names):
        hdr_cells[i].text = col_name

    # Add rows
    for row in data:
        row_cells = table.add_row().cells
        for i, col_name in enumerate(column_names):
            row_cells[i].text = str(row.get(col_name, ""))

def add_table_of_contents(doc, content=None, images=None, tables=None):
    """Add a table of contents with proper Word field references"""
    # Add Table of Contents heading
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add some spacing
    doc.add_paragraph()
    
    # Define the sections that will be in the document
    sections = [
        "Attack Indicators",
        "Honeypot Attack Trends", 
        "Network Traffic by Protocol",
        "Indicator of Attacks",
        "Top IP Addresses",
        "Credential Patterns",
        "Subdomains",
        "Email Addresses",
        "Hashes",
        "About Sequretek"
    ]
    
    # Create TOC entries with PAGEREF fields
    for section_name in sections:
        toc_para = doc.add_paragraph()
        
        # Set up tab stops for page numbers (right-aligned with dots)
        toc_para.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # Add section name
        toc_para.add_run(section_name)
        
        # Add tab
        toc_para.add_run("\t")
        
        # Create bookmark name (same as in add_heading)
        bookmark_name = section_name.replace(' ', '_').replace('&', 'and').replace(':', '').replace('/', '')
        
        # Add PAGEREF field that references the bookmark
        # Create the PAGEREF field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = f'PAGEREF {bookmark_name} \\h'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        # Add the field to a new run
        field_run = toc_para.add_run()
        field_run._r.append(fldChar1)
        field_run._r.append(instrText)
        field_run._r.append(fldChar2)
        
        # Add result text (will be updated when document is opened in Word)
        result_text = OxmlElement('w:t')
        result_text.text = "1"  # Placeholder that Word will replace
        field_run._r.append(result_text)
    
    # Add page break after TOC
    doc.add_page_break()

def add_footer(doc, section=None):
    """Add a footer to the document that appears on every page"""
    if section is None:
        section = doc.sections[0]
    
    footer = section.footer
    
    # Clear any existing footer content
    for para in footer.paragraphs:
        para.clear()
    
    # Remove all existing paragraphs except the first one
    while len(footer.paragraphs) > 1:
        footer._element.remove(footer.paragraphs[1]._element)
    
    # Use a single paragraph with left alignment
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Configure paragraph format - clear any existing tab stops
    pformat = footer_para.paragraph_format
    pformat.tab_stops.clear_all()
    
    # Clear and add content
    footer_para.clear()
    
    # Only "©Sequretek" text on the left
    left_run = footer_para.add_run("©Sequretek")
    left_run.font.name = 'Calibri'
    left_run.font.size = Pt(10)

def add_header(doc, section=None):
    """Add a header to the document that appears on every page"""
    if section is None:
        section = doc.sections[0]
    
    header = section.header
    
    # Clear any existing header content
    header.paragraphs[0].clear()
    
    # Create a single paragraph for the header
    header_para = header.paragraphs[0]
    
    # Set up tab stops for proper alignment
    header_para.paragraph_format.tab_stops.clear_all()
    header_para.paragraph_format.tab_stops.add_tab_stop(Inches(8.5), WD_TAB_ALIGNMENT.RIGHT)
    
    # Left side - Page number
    header_para.add_run("Page ")
    
    # Add page number field with proper formatting
    run = header_para.runs[-1]
    
    # Create the PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = " PAGE \\* Arabic "  # Use Arabic numerals
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    # Add the field to the run
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # Right side - Always use Sequretek logo from templates folder
    sequretek_logo_path = "templates/sequretek.png"
    try:
        # Add tab to move to right side
        header_para.add_run("\t")
        # Create a new paragraph for the logo to avoid conflicts
        logo_para = header.add_paragraph()
        logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        logo_run = logo_para.add_run()
        logo_run.add_picture(sequretek_logo_path, width=Inches(1.5))
    except Exception as e:
        # If logo fails to load, add text placeholder
        header_para.add_run("\tSequretek Labs")
    
    # Set paragraph alignment
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_cover_page(doc, org_name=None, org_logo_path=None, analysis_start=None, analysis_end=None):
    """Add a professional cover page to the document with two-tone background"""
    # Ensure there's at least one paragraph to work with
    if len(doc.paragraphs) == 0:
        doc.add_paragraph()
    
    # Clear the first paragraph
    doc.paragraphs[0].clear()
    
    # Set page margins to be smaller
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Add very minimal top spacing and Sequretek logo at the top right
    sequretek_logo_path = "templates/sequretek.png"
    try:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        logo_para.space_before = Pt(6)  # Very small top margin
        logo_para.space_after = Pt(6)   # Very small bottom margin
        logo_run = logo_para.add_run()
        logo_run.add_picture(sequretek_logo_path, width=Inches(2))
    except Exception:
        # If logo fails, add text
        logo_para = doc.add_paragraph("SEQURETEK")
        logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        logo_para.space_before = Pt(6)
        logo_para.space_after = Pt(6)
        logo_run = logo_para.runs[0]
        logo_run.font.size = Pt(18)
        logo_run.font.bold = True
        logo_run.font.color.rgb = RGBColor(255, 255, 255)  # White text on dark background
    
    # Create dark blue background section (top half)
    # Add a table that spans the width to create background effect
    bg_table = doc.add_table(rows=1, cols=1)
    bg_table.style = 'Table Grid'
    bg_cell = bg_table.rows[0].cells[0]
    
    # Set cell background to dark blue
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
    shading_elm = parse_xml(r'<w:shd {} w:fill="1B1B70"/>'.format(nsdecls('w')))
    bg_cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # Set cell padding and height
    bg_cell.width = Inches(7)
    
    # Add content to the dark blue section
    bg_para = bg_cell.paragraphs[0]
    bg_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    bg_para.space_before = Pt(12)
    bg_para.space_after = Pt(12)
    
    # Add main title "DECOY REPORT" in white text
    title_run = bg_para.add_run("DECOY REPORT")
    title_run.font.size = Pt(42)  # Larger title
    title_run.font.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Add line break and "Prepared by: Sequretek Lab" in white
    bg_para.add_run("\n\n")
    prepared_run = bg_para.add_run("Prepared by: ")
    prepared_run.font.size = Pt(16)
    prepared_run.font.name = 'Calibri'
    prepared_run.font.color.rgb = RGBColor(255, 255, 255)
    
    sequretek_run = bg_para.add_run("Sequretek Lab")
    sequretek_run.font.size = Pt(16)
    sequretek_run.font.bold = True
    sequretek_run.font.name = 'Calibri'
    sequretek_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Add decorative line
    bg_para.add_run("\n")
    underline_run = bg_para.add_run("_" * 35)
    underline_run.font.color.rgb = RGBColor(255, 255, 255)
    underline_run.font.size = Pt(14)
    
    # Minimal spacing after dark section
    spacing_para = doc.add_paragraph()
    spacing_para.space_before = Pt(6)
    spacing_para.space_after = Pt(6)
    
    # Add the cybersecurity analyst image (larger)
    try:
        image_para = doc.add_paragraph()
        image_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        image_para.space_before = Pt(6)  # Reduced spacing
        image_para.space_after = Pt(6)   # Reduced spacing
        image_run = image_para.add_run()
        image_run.add_picture("templates/base.png", width=Inches(5.5))  # Larger image
    except Exception:
        # If image fails to load, add placeholder text
        image_para = doc.add_paragraph()
        image_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        image_para.space_before = Pt(6)
        image_para.space_after = Pt(6)
        image_run = image_para.add_run("[Cybersecurity Operations Center Image]")
        image_run.font.italic = True
        image_run.font.size = Pt(14)
        image_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Minimal spacing
    spacing_para2 = doc.add_paragraph()
    spacing_para2.space_before = Pt(6)
    spacing_para2.space_after = Pt(6)
    
    # Add Organization Logo/Image (from logos folder)
    if org_logo_path:
        try:
            org_logo_para = doc.add_paragraph()
            org_logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            org_logo_para.space_before = Pt(6)
            org_logo_para.space_after = Pt(6)
            org_logo_run = org_logo_para.add_run()
            org_logo_run.add_picture(org_logo_path, width=Inches(1.8))  # Slightly larger org logo
        except Exception:
            # If org logo fails, add placeholder text
            org_logo_para = doc.add_paragraph()
            org_logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            org_logo_para.space_before = Pt(6)
            org_logo_para.space_after = Pt(6)
            org_logo_run = org_logo_para.add_run("Organization Logo/Image")
            org_logo_run.font.size = Pt(16)
            org_logo_run.font.bold = True
            org_logo_run.font.name = 'Calibri'
    else:
        # No org logo provided, add placeholder
        org_logo_para = doc.add_paragraph()
        org_logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        org_logo_para.space_before = Pt(6)
        org_logo_para.space_after = Pt(6)
        org_logo_run = org_logo_para.add_run("Organization Logo/Image")
        org_logo_run.font.size = Pt(16)
        org_logo_run.font.bold = True
        org_logo_run.font.name = 'Calibri'
    
    # Add Organization Name (from logos folder)
    org_name_para = doc.add_paragraph()
    org_name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    org_name_para.space_before = Pt(6)
    org_name_para.space_after = Pt(6)
    if org_name:
        org_name_run = org_name_para.add_run(org_name)
    else:
        org_name_run = org_name_para.add_run("Organization Name")
    org_name_run.font.size = Pt(20)  # Slightly larger
    org_name_run.font.bold = True
    org_name_run.font.name = 'Calibri'
    org_name_run.font.color.rgb = RGBColor(25, 25, 112)  # Dark blue text
    
    # Use provided dates or fall back to current date
    if analysis_start and analysis_end:
        analysis_start_str = analysis_start.strftime("%d/%m/%Y")
        analysis_end_str = analysis_end.strftime("%d/%m/%Y")
        issued_date_str = datetime.now().strftime("%d/%m/%Y")
    else:
        # Fallback to current date information
        current_date = datetime.now()
        analysis_start_str = current_date.replace(day=1).strftime("%d/%m/%Y")
        analysis_end_str = current_date.strftime("%d/%m/%Y")
        issued_date_str = current_date.strftime("%d/%m/%Y")
    
    # Analysis Period and Issued Date (compact)
    analysis_para = doc.add_paragraph()
    analysis_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    analysis_para.space_before = Pt(6)
    analysis_para.space_after = Pt(6)
    analysis_run = analysis_para.add_run(f"Analysis Period: {analysis_start_str} to {analysis_end_str}")
    analysis_run.font.size = Pt(11)  # Smaller text
    analysis_run.font.name = 'Calibri'
    
    # Issued Date (same line to save space)
    issued_run = analysis_para.add_run(f" | Issued Date: {issued_date_str}")
    issued_run.font.size = Pt(11)  # Smaller text
    issued_run.font.name = 'Calibri'
    
    # Add page break after cover page
    doc.add_page_break()

def add_about_sequretek_page(doc):
    """Add the About Sequretek page at the end of the report"""
    # Add page break before the about page
    doc.add_page_break()
    
    # Add the title
    about_heading = doc.add_heading("About Sequretek", level=1)
    about_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add some spacing after title
    spacing_para = doc.add_paragraph()
    spacing_para.space_before = Pt(12)
    spacing_para.space_after = Pt(12)
    
    # Add the first paragraph
    para1 = doc.add_paragraph()
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para1_text = ("Built on the Percept CTEM framework, Sequretek's products—Percept XDR & NG SIEM, "
                  "Identity, EDR, and Compliance Manager—offer defense-in-depth and defense-in-breadth "
                  "capabilities. Our AI technologies, processes, and expert teams ensure organizations' "
                  "IT assets are protected against cyberthreats, while helping them comply with regulatory "
                  "standards. Our unwavering commitment to provide effortless cybersecurity with cost-effective "
                  "solutions, empowers our customers to grow confidently while navigating the digital world.")
    
    para1_run = para1.add_run(para1_text)
    para1_run.font.size = Pt(12)
    para1_run.font.name = 'Calibri'
    
    # Add spacing between paragraphs
    spacing_para2 = doc.add_paragraph()
    spacing_para2.space_before = Pt(6)
    spacing_para2.space_after = Pt(6)
    
    # Add the second paragraph
    para2 = doc.add_paragraph()
    para2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para2_text = ("Sequretek is recognized by leading global research & analyst firms. We are a winner "
                  "of the TiE50 awards, National Startup Awards, NASSCOM Emerge 50 awards, and made it "
                  "to the Financial Times High Growth Asia Pacific companies. Sequretek has featured in "
                  "the list of Top 250 MSSPs across the globe by MSSP Alert for three years in a row.")
    
    para2_run = para2.add_run(para2_text)
    para2_run.font.size = Pt(12)
    para2_run.font.name = 'Calibri'
    
    # Add some spacing at the bottom
    spacing_para3 = doc.add_paragraph()
    spacing_para3.space_before = Pt(18)
    spacing_para3.space_after = Pt(18)
    
    # Add Sequretek logo at the bottom center (optional)
    try:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture("templates/sequretek.png", width=Inches(2.5))
    except Exception:
        # If logo fails, add company name
        logo_para = doc.add_paragraph("SEQURETEK")
        logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        logo_run = logo_para.runs[0]
        logo_run.font.size = Pt(20)
        logo_run.font.bold = True
        logo_run.font.name = 'Calibri'
        logo_run.font.color.rgb = RGBColor(25, 25, 112)

def generate_docx_report(content: dict, images: dict, tables: dict, output_path: str, analysis_dates: tuple = None, logo_folder_path: str = None):
    doc = Document()
    
    # Extract analysis dates if provided
    analysis_start = None
    analysis_end = None
    if analysis_dates:
        analysis_start, analysis_end = analysis_dates
    
    org_name, org_logo_path = get_organization_info(logo_folder_path)
    
    # Add cover page first (this section will have no header/footer)
    add_cover_page(doc, org_name, org_logo_path, analysis_start, analysis_end)
    
    # Create a new section for the rest of the document (TOC and content)
    new_section = doc.add_section()
    
    # Configure the new section
    new_section.start_type = 2  # New page
    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False
    
    # Set up page numbering for the new section
    from docx.oxml.ns import qn
    sectPr = new_section._sectPr
    
    # Remove any existing page number type
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        sectPr.remove(pgNumType)
    
    # Create new page number type with explicit settings
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')  # Start from 1
    pgNumType.set(qn('w:fmt'), 'decimal')  # Use decimal numbers
    pgNumType.set(qn('w:chapStyle'), '0')  # No chapter style
    pgNumType.set(qn('w:chapSep'), 'none')  # No chapter separator
    sectPr.append(pgNumType)
    
    # Add header and footer to the new section
    add_header(doc, section=new_section)
    add_footer(doc, new_section)

    # Add Table of Contents
    add_table_of_contents(doc)

    # Add content sections
    for section_name in [
        "Attack Indicators",
        "Honeypot Attack Trends",
        "Network Traffic by Protocol",
        "Indicator of Attacks",
        "Top IP Addresses",
        "Credential Patterns",
        "Subdomains",
        "Email Addresses",
        "Hashes",
    ]:
        text = content.get(section_name)
        chart = images.get(section_name)
        table = tables.get(section_name)

        add_heading(doc, section_name, level=1)

        if chart:
            add_image(doc, chart)

        if table:
            # Handle both single table and list of tables
            if isinstance(table, list):
                for single_table in table:
                    if single_table and len(single_table) > 0:
                        column_names = list(single_table[0].keys())
                        add_table(doc, single_table, column_names)
            elif len(table) > 0:
                column_names = list(table[0].keys())
                add_table(doc, table, column_names)

        if text:
            add_paragraph(doc, text)

    # Add the About Sequretek page at the end
    add_about_sequretek_page(doc)

    doc.save(output_path)
    print(f"✅ Final report saved to {output_path}")